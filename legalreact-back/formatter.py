import os
import json
import re
from azure.storage.blob import BlobServiceClient
from langchain_openai import AzureChatOpenAI
from docx import Document
from difflib import get_close_matches

# 🔹 Azure Blob Storage Configuration
AZURE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
CONTAINER_NAME = os.getenv("AZURE_CONTAINER_NAME_4")

# 🔹 OpenAI Configuration
OPENAI_API_KEY = os.getenv("OPENAI_GPT_API_KEY")
OPENAI_API_ENDPOINT = os.getenv("OPENAI_GPT_ENDPOINT")

llm = AzureChatOpenAI(
    azure_deployment="gpt-4o-mini",
    azure_endpoint=OPENAI_API_ENDPOINT,
    api_key=OPENAI_API_KEY,
    api_version="2024-10-21",
    temperature=0.2
)

# 🔹 Function to list available templates from Azure Blob Storage
def list_templates_from_blob():
    try:
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_CONNECTION_STRING)
        container_client = blob_service_client.get_container_client(CONTAINER_NAME)
        template_files = [blob.name for blob in container_client.list_blobs() if blob.name.endswith(".docx")]
        template_names = [os.path.splitext(os.path.basename(file))[0] for file in template_files]
        return template_names
    except Exception as e:
        print(f"Error listing templates: {e}")
        return []


# 🔹 Function to classify document type using GPT
def classify_document_type(user_query):
    try:
        response = llm.invoke([
            {"role": "system", "content": "Identify the most appropriate document type that is being asked for in the user query. Only return the document type name with no extra text. If it is a non disclosure agreement then return NDA. Valid template names should have NDA or Business Partnership in them."},
            {"role": "user", "content": user_query}
        ])

        print(f"Raw LLM Response: {response.content.strip()}")  # Debugging

        if response and hasattr(response, 'content'):
            document_type = response.content.strip().upper().replace(" ", "_")
            print(f"Classified Document Type: {document_type}")  # Debugging
            return document_type
        print("Error: Classification failed.")
        return None
    except Exception as e:
        print(f"Error in document classification: {e}")
        return None


# 🔹 Function to fetch template from Azure Blob Storage
def fetch_template_from_blob(document_type):
    try:
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_CONNECTION_STRING)
        container_client = blob_service_client.get_container_client(CONTAINER_NAME)

        # Extract available template names
        available_templates = [blob.name.replace(".docx", "") for blob in container_client.list_blobs()]
        print(f"Available templates in container: {available_templates}")  # Debugging

        # Find the closest match
        best_match = get_close_matches(document_type, available_templates, n=1, cutoff=0.5)  # Adjust cutoff if needed

        if not best_match:
            print(f"Error: No matching template found for {document_type}")
            return None

        selected_template = best_match[0]  # Get the best match

        # Download the matched template
        blob_client = blob_service_client.get_blob_client(container=CONTAINER_NAME, blob=f"{selected_template}.docx")
        template_path = f"templates/{selected_template}.docx"
        os.makedirs("templates", exist_ok=True)  # Ensure directory exists

        with open(template_path, "wb") as template_file:
            template_file.write(blob_client.download_blob().readall())

        print(f"Template {selected_template}.docx downloaded successfully.")  # Debugging
        return template_path
    except Exception as e:
        print(f"Error fetching template: {e}")
        return None


# 🔹 Function to extract placeholders from the document template
def extract_placeholders(template_path):
    try:
        doc = Document(template_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        placeholders = list(set(re.findall(r"\{(.*?)\}", text)))  # Extract placeholders
        return placeholders
    except Exception as e:
        print(f"Error extracting placeholders: {e}")
        return []


# 🔹 Function to generate the extraction prompt for GPT
def generate_extraction_prompt(user_query, document_type, placeholders):
    """
    Generate the extraction prompt for GPT.
    """
    extraction_prompt = f"""

    **User Description:**  
    The input is {user_query}. Please extract the required details from it.
    
    1. **Identify the parties involved**:
    - Look for names of individuals or companies in the description.
    - Classify them into:
        - **DISCLOSING_PARTIES** (e.g., the company or primary entity sharing information)
        - **RECEIVING_PARTIES** (e.g., the partner, recipient, or secondary entity)
    - Ensure that you list the parties in a **separate list** under each category. For example:
        - "DISCLOSING_PARTIES": ["Company X", "Company Y"]
        - "RECEIVING_PARTIES": ["Person A", "Person B"]]

    The following is a description of a {document_type}. Please extract the required details from it.

    2. **Identify and extract the following placeholders**:
    - **AGREEMENT_DATE**: The date the agreement is made, in **YYYY-MM-DD** format.
    - **COMMENCEMENT_DATE**: The date the agreement becomes effective, in **YYYY-MM-DD** format.
    - **TERM_YEARS**: The term of the agreement, in **years** (e.g., "5").
    - Ensure you populate these values based on the information in the description.

    3. **Return the data strictly in JSON format** with the following structure:
    - If any information is missing or unclear, leave the respective field as `null`.
    
    ---

    **Expected Output Format (Strict JSON)**:
    {{
        "AGREEMENT_DATE": "YYYY-MM-DD",
        "COMMENCEMENT_DATE": "YYYY-MM-DD",
        "TERM_YEARS": "X",
        "DISCLOSING_PARTIES": ["Name1", "Name2"],
        "RECEIVING_PARTIES": ["Name1", "Name2"]
    }}
    """
    return extraction_prompt


# 🔹 Function to extract entities from GPT response
def extract_entities_from_gpt(user_query, document_type, placeholders):
    """
    Calls GPT to extract entities from the user query.
    """
    extraction_prompt = generate_extraction_prompt(user_query, document_type, placeholders)
    
    # Call the LLM with the enhanced prompt
    response = llm.invoke([{"role": "user", "content": extraction_prompt}])
    extracted_data = extract_json_from_response(response.content.strip())
    
    return extracted_data


def extract_json_from_response(response_text):
    try:
        # Remove markdown formatting if present (e.g., ```json ... ```)
        response_text = re.sub(r"```json\n|\n```", "", response_text).strip()
        print(f"GPT Response: {response_text}")
        # Convert to JSON
        extracted_data = json.loads(response_text)

        return extracted_data

    except json.JSONDecodeError:
        return None  # Return None if parsing fails


def fill_document_with_gpt(template_path, extracted_data):
    try:
        doc = Document(template_path)

        # Ensure placeholders are properly formatted
        cleaned_data = {
            key: (str(value) if value is not None else "") for key, value in extracted_data.items()
        }

        # Convert list elements into properly formatted strings
        for key in ["DISCLOSING_PARTIES", "RECEIVING_PARTIES"]:
            if key in extracted_data and isinstance(extracted_data[key], list):
                cleaned_data[key] = ", ".join(map(str, extracted_data[key]))  # Join correctly

        # Replace placeholders with extracted values
        for para in doc.paragraphs:
            for placeholder, value in cleaned_data.items():
                para.text = para.text.replace(f"{{ {placeholder} }}", value)

        # Replace custom Jinja-style placeholders
        for para in doc.paragraphs:
            if "{% for party in DISCLOSING_PARTIES %}" in para.text:
                para.text = para.text.replace(
                    "{% for party in DISCLOSING_PARTIES %} - { party.name }", 
                    f"Party Disclosing Information: {cleaned_data.get('DISCLOSING_PARTIES', 'N/A')}"
                )

            if "{% for party in RECEIVING_PARTIES %}" in para.text:
                para.text = para.text.replace(
                    "{% for party in RECEIVING_PARTIES %} - { party.name }", 
                    f"Party Receiving Information: {cleaned_data.get('RECEIVING_PARTIES', 'N/A')}"
                )

        # Save the final document
        output_path = "/Users/aryan_zingade/Downloads/generated_document.docx"
        doc.save(output_path)
        return output_path

    except Exception as e:
        print(f"Error in document generation: {e}")
        return None



